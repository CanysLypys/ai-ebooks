import json
import datetime
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

client = OpenAI(
    api_key="Insert API Key"
)

def generate_ebook_content(topic, target_audience, num_chapters, num_subsections):
    prompt = (
        f'We are writing an eBook titled "{topic}" targeted at "{target_audience}".'
        f' Provide a comprehensive outline for the eBook with {num_chapters} chapter(s).'
        f' Each chapter should have exactly {num_subsections} subsection(s) related to the topic of {topic}.'
        'Output format for prompt:'
        ' python dict with key: chapter title, value: a single list/array'
        ' containing subsection titles within the chapter (the subtopics'
        ' should be inside the list).'
    )

    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ],
        model="gpt-3.5-turbo",
    )

    outline = json.loads(chat_completion.choices[0].message.content.strip())

    doc = Document()

    title_page = doc.add_paragraph(topic)
    title_page.style.font.size = Pt(12)
    title_page.style.font.bold = True
    title_page.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for chapter_title, _ in outline.items():
        toc_entry = doc.add_paragraph(chapter_title)
        toc_entry.style.font.size = Pt(14)
        toc_entry.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        toc_entry.add_run('\n')

    doc.add_page_break()

    for chapter_title, subsections in outline.items():
        doc.add_heading(chapter_title, level=1)

        for subsection_title in subsections:
            subsection_prompt = f"Write a comprehensive and detailed content for the subsection on {subsection_title}. Please include complete information about {subsection_title} and ensure that the content is thorough and well-explained. Keep in mind that the overall topic is {chapter_title}."
            subsection_completion = client.chat.completions.create(
                messages=[
                    {
                        "role": "user",
                        "content": subsection_prompt
                    }
                ],
                model="gpt-3.5-turbo"
            )

            doc.add_paragraph(subsection_completion.choices[0].message.content.strip())
            doc.add_page_break()

    current_datetime = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    ebook_name = f"E_Book_{current_datetime}.docx"
    doc.save(ebook_name)

    return ebook_name

def main():
    topic = input("Enter the topic of the eBook: ")
    target_audience = input("Enter the target audience of the eBook: ")
    num_chapters = int(input("Enter the number of chapters: "))
    num_subsections = int(input("Enter the number of subsections per chapter: "))

    print("Generating eBook...")

    e_book_name = generate_ebook_content(topic, target_audience, num_chapters, num_subsections)

    print(f"EBook generated successfully. File saved as: {e_book_name}")

if __name__ == "__main__":
    main()
